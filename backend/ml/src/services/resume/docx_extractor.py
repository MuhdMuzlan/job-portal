"""DOCX text extractor"""

import logging
from typing import Optional
from docx import Document
import io

logger = logging.getLogger(__name__)


class DocxExtractor:
    """Extract text from DOCX files"""
    
    async def extract(self, content: bytes) -> str:
        """
        Extract text from DOCX content
        
        Args:
            content: DOCX file bytes
            
        Returns:
            Extracted text
        """
        try:
            doc = Document(io.BytesIO(content))
            
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            full_text = '\n\n'.join(text_parts)
            
            return full_text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            raise
